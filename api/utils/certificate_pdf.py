from io import BytesIO
from django.core.files.base import ContentFile
from datetime import date
from docx import Document
# from weasyprint import HTML
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from io import BytesIO
from PyPDF2 import PdfReader, PdfWriter
import qrcode
import base64
from docx.shared import Inches
from docx.shared import Pt
from PIL import Image
import os


def generate_certificate(student, course, template_obj,certificate_id):
    if template_obj.file_type == 'html':
        return generate_from_html(student, course, template_obj.html_template,certificate_id)
    elif template_obj.file_type == 'docx':
        return generate_from_docx(student, course, template_obj.file,certificate_id)
    # elif template_obj.file_type == 'pdf':
    #     return generate_from_pdf(student, course, template_obj.file)
    else:
        raise ValueError("Unsupported template type")


def get_name(obj):
    user = getattr(obj, 'user', obj)  # Works for both Student and Instructor
    name = f"{user.first_name or ''} {user.last_name or ''}".strip()
    return name if name else user.username



def generate_from_html(student, course, html_template, certificate_id):
   
    # 1. Create QR code
    qr_url = f"http://127.0.0.1:8000/api/certificate/verify/{certificate_id}"
    qr_img = qrcode.make(qr_url)

    # 2. Resize to smaller dimensions (e.g., 100x100 px)
    qr_img = qr_img.resize((100, 100), Image.LANCZOS)

    # 3. Convert to base64
    qr_io = BytesIO()
    qr_img.save(qr_io, format='PNG')
    qr_io.seek(0)
    qr_base64 = base64.b64encode(qr_io.read()).decode('utf-8')

  # 4. Escape all curly braces
    safe_html = html_template.replace("{", "{{").replace("}", "}}")

    # 5. Replace placeholders safely
    safe_html = (
        safe_html
        .replace("{{student_name}}", get_name(student))
        .replace("{{course_title}}", course.title)
        .replace("{{instructor_name}}", get_name(course.instructor))
        .replace("{{date}}", str(date.today()))
    )

    # 5. Inject QR and ID at the bottom
    qr_html = f"""
        <div style="margin-top: 40px; text-align: right;">
            <p style="font-size: 12px;"><strong>Certificate ID:</strong> {certificate_id}</p>
            <img src="data:image/png;base64,{qr_base64}" alt="QR Code" width="100" height="100" />
            <p style="font-size: 10px; color: #666;">Scan to verify certificate</p>
        </div>
    """

    final_html = safe_html + qr_html

    # 6. Generate and return PDF
    pdf_file = BytesIO()
    HTML(string=final_html).write_pdf(pdf_file)
    pdf_file.seek(0)
    return ContentFile(pdf_file.read(), name=f"certificate_{student.id}_{course.id}.pdf")



def generate_from_docx(student, course, docx_file, certificate_id):
    doc = Document(docx_file)
    
    fields = {
        '{{student_name}}': get_name(student),
        '{{course_title}}': course.title,
        '{{instructor_name}}': get_name(course.instructor),
        '{{date}}': str(date.today()),
        '{{certificate_id}}': certificate_id
    }

    # Replace text placeholders
    for para in doc.paragraphs:
        for key, value in fields.items():
            if key in para.text:
                para.text = para.text.replace(key, value)

    # Generate QR code
    qr_url = f"http://127.0.0.1:8000/api/certificate/verify/{certificate_id}"
    qr_img = qrcode.make(qr_url)
    qr_path = f"temp_qr_{certificate_id}.png"
    qr_img.save(qr_path)

    # Insert at the end
    doc.add_paragraph(" ")
    doc.add_paragraph("Certificate ID: " + certificate_id)
    doc.add_picture(qr_path, width=Inches(1.2))

    # Save and clean
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    os.remove(qr_path)

    return ContentFile(output.read(), name=f"certificate_{student.id}_{course.id}.docx")



# def generate_from_pdf(student, course, pdf_template_file):

#     # Step 1: Load template
#     template_reader = PdfReader(pdf_template_file)
#     template_page = template_reader.pages[0]

#     # Step 2: Replace fields visually (you need to manually match the positions)
#     overlay_packet = BytesIO()
#     can = canvas.Canvas(overlay_packet, pagesize=letter)

#     # Values to replace
#     fields = {
#         'student_name': get_name(student),
#         'course_title': course.title,
#         'instructor_name': get_name(course.instructor),
#         'date': str(date.today())
#     }

#     # Step 3: Overlay text at the correct (x, y) positions
#     # You must find these manually by testing or using Adobe Acrobat to inspect the layout

#     can.setFont("Helvetica-Bold", 20)
#     can.drawString(180, 505, fields['student_name'])     # Replace {student_name}
#     can.drawString(180, 465, fields['course_title'])     # Replace {course_title}

#     can.setFont("Helvetica", 14)
#     can.drawString(130, 430, f"Instructor: {fields['instructor_name']}")  # Replace {instructor_name}
#     can.drawString(130, 400, f"Date: {fields['date']}")                   # Replace {date}

#     can.save()
#     overlay_packet.seek(0)

#     # Step 4: Merge overlay with template
#     overlay_pdf = PdfReader(overlay_packet)
#     output = PdfWriter()

#     # Merge only the first page
#     template_page.merge_page(overlay_pdf.pages[0])
#     output.add_page(template_page)

#     final_output = BytesIO()
#     output.write(final_output)
#     final_output.seek(0)

#     return ContentFile(final_output.read(), name=f"certificate_{student.id}_{course.id}.pdf")
